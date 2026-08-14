"""
Unit tests for PII detectors, faker replacement consistency, and evaluation metrics.
"""

import pytest
import docx
import tempfile
import os
import json
from src.detectors import detect_pii, is_luhn_valid, validate_ip
from src.redactor import PIIRedactor
from src.evaluator import evaluate_redaction


def test_email_detection():
    text = "Contact us at john.doe@example.com for inquiries."
    detections = detect_pii(text)
    email_dets = [d for d in detections if d['type'] == 'EMAIL']
    assert len(email_dets) == 1
    assert email_dets[0]['text'] == "john.doe@example.com"


def test_phone_detection():
    text = "Call our hotline at +91 9876543210 or 020 4505 3237."
    detections = detect_pii(text)
    phone_dets = [d for d in detections if d['type'] == 'PHONE']
    assert len(phone_dets) >= 1
    assert any("9876543210" in d['text'] for d in phone_dets)


def test_ssn_detection():
    text = "Social Security Number: 123-45-6789."
    detections = detect_pii(text)
    ssn_dets = [d for d in detections if d['type'] == 'SSN']
    assert len(ssn_dets) == 1
    assert ssn_dets[0]['text'] == "123-45-6789"


def test_credit_card_detection():
    valid_card = "4111 1111 1111 1111"
    text = f"Payment method: {valid_card}"
    assert is_luhn_valid(valid_card)
    detections = detect_pii(text)
    card_dets = [d for d in detections if d['type'] == 'CREDIT_CARD']
    assert len(card_dets) == 1


def test_invalid_credit_card_rejected():
    invalid_card = "1111 1111 1111 1111"
    text = f"Invalid card: {invalid_card}"
    assert not is_luhn_valid(invalid_card)
    detections = detect_pii(text)
    card_dets = [d for d in detections if d['type'] == 'CREDIT_CARD']
    assert len(card_dets) == 0


def test_ip_address_detection():
    text = "Server IP: 192.168.1.10"
    detections = detect_pii(text)
    ip_dets = [d for d in detections if d['type'] == 'IP_ADDRESS']
    assert len(ip_dets) == 1
    assert ip_dets[0]['text'] == "192.168.1.10"


def test_invalid_ip_address_rejected():
    assert not validate_ip("999.999.999.999")
    text = "Invalid IP: 999.999.999.999"
    detections = detect_pii(text)
    ip_dets = [d for d in detections if d['type'] == 'IP_ADDRESS']
    assert len(ip_dets) == 0


def test_dob_detection():
    text = "Candidate Date of Birth: 15/08/1999."
    detections = detect_pii(text)
    dob_dets = [d for d in detections if d['type'] == 'DOB']
    assert len(dob_dets) == 1
    assert dob_dets[0]['text'] == "15/08/1999"


def test_ordinary_date_not_dob():
    text = "The prospectus was filed on December 10, 2025."
    detections = detect_pii(text)
    dob_dets = [d for d in detections if d['type'] == 'DOB']
    assert len(dob_dets) == 0


def test_person_detection():
    text = "The Managing Director is Rashi Patil."
    detections = detect_pii(text)
    person_dets = [d for d in detections if d['type'] == 'PERSON']
    assert len(person_dets) == 1
    assert "Rashi Patil" in person_dets[0]['text']


def test_organization_detection():
    text = "The audit was conducted by Example Technologies Private Limited."
    detections = detect_pii(text)
    org_dets = [d for d in detections if d['type'] == 'ORGANIZATION']
    assert len(org_dets) == 1
    assert "Example Technologies" in org_dets[0]['text']


def test_non_pii_order_id_not_redacted():
    text = "Order ID: 123456 with 50 shares."
    detections = detect_pii(text)
    assert len(detections) == 0


def test_replacement_consistency():
    redactor = PIIRedactor(seed=42)
    rep1 = redactor.get_replacement("Rashi Patil", "PERSON")
    rep2 = redactor.get_replacement("Rashi Patil", "PERSON")
    assert rep1 == rep2
    assert rep1 != "Rashi Patil"


def test_replacement_case_normalization():
    redactor = PIIRedactor(seed=42)
    rep_title = redactor.get_replacement("Rashi Patil", "PERSON")
    rep_upper = redactor.get_replacement("RASHI PATIL", "PERSON")
    assert rep_title.upper() == rep_upper


def test_distinct_phones_have_different_replacements():
    redactor = PIIRedactor(seed=42)
    phone1_rep = redactor.get_replacement("+91 9876543210", "PHONE")
    phone2_rep = redactor.get_replacement("+91 20 45053237", "PHONE")
    assert phone1_rep != phone2_rep


def test_label_based_name_detection():
    """Names following 'Full Name:' or 'Name:' labels should be detected."""
    text = "Full Name: Rahul Sharma\nName: Priya Mehta"
    detections = detect_pii(text)
    person_texts = [d['text'] for d in detections if d['type'] == 'PERSON']
    assert any("Rahul Sharma" in t for t in person_texts), "Rahul Sharma not detected"
    assert any("Priya Mehta" in t for t in person_texts), "Priya Mehta not detected"


def test_label_based_organization_detection():
    """Company names following 'Company:' label should be detected."""
    text = "Company:\nSharma Technologies Private Limited"
    detections = detect_pii(text)
    org_texts = [d['text'] for d in detections if d['type'] == 'ORGANIZATION']
    assert any("Sharma Technologies" in t for t in org_texts), "Sharma Technologies not detected"


def test_label_based_address_detection():
    """Addresses with a 6-digit Indian PIN code following 'Address:' should be detected."""
    text = "Address:\n42 Green Park Road, Sector 18,\nNoida, Uttar Pradesh 201301, India"
    detections = detect_pii(text)
    addr_texts = [d['text'] for d in detections if d['type'] == 'ADDRESS']
    assert len(addr_texts) >= 1, "Address not detected"
    assert any("201301" in t for t in addr_texts), "PIN code not found in detected address"


def test_credit_card_with_spaces():
    """Credit card with spaces should pass Luhn and be detected."""
    text = "Credit Card: 4111 1111 1111 1111"
    assert is_luhn_valid("4111 1111 1111 1111")
    detections = detect_pii(text)
    card_dets = [d for d in detections if d['type'] == 'CREDIT_CARD']
    assert len(card_dets) == 1, f"Credit card not detected. Got: {detections}"


def test_all_9_categories_synthetic():
    """Test that all 9 PII categories are caught in a synthetic test string."""
    text = (
        "Full Name: Rahul Sharma\n"
        "Email: rahul.sharma@example.com\n"
        "Phone: +91 9876543210\n"
        "Date of Birth: 15/08/2001\n"
        "Address:\n"
        "42 Green Park Road, Sector 18,\n"
        "Noida, Uttar Pradesh 201301, India\n"
        "Company:\n"
        "Sharma Technologies Private Limited\n"
        "SSN: 123-45-6789\n"
        "Credit Card: 4111 1111 1111 1111\n"
        "IP Address: 192.168.1.25\n"
    )
    detections = detect_pii(text)
    found_types = {d['type'] for d in detections}
    required = {"PERSON", "EMAIL", "PHONE", "DOB", "ADDRESS", "ORGANIZATION", "SSN", "CREDIT_CARD", "IP_ADDRESS"}
    missing = required - found_types
    assert not missing, f"Missing categories: {missing}"


def test_name_label_does_not_cross_newline_into_next_label():
    """Person name pattern should not spill over into the next line's label."""
    text = "Full Name: Aarav Mehta\nEmail: aarav.mehta@example.com"
    detections = detect_pii(text)
    person_texts = [d['text'] for d in detections if d['type'] == 'PERSON']
    assert any(t == "Aarav Mehta" for t in person_texts), f"Got: {person_texts}"
    assert not any("Email" in t for t in person_texts), f"Name crossed into next label: {person_texts}"


def test_address_does_not_swallow_next_section_label():
    """Address pattern should stop before the next section label."""
    text = (
        "Address:\n"
        "28 MG Road, Andheri West,\n"
        "Mumbai, Maharashtra 400058, India\n"
        "Credit Card: 5555 5555 5555 4444"
    )
    detections = detect_pii(text)
    addr_texts = [d['text'] for d in detections if d['type'] == 'ADDRESS']
    assert len(addr_texts) >= 1, "Address not detected"
    assert all("Credit Card" not in t for t in addr_texts), f"Address swallowed next section: {addr_texts}"


def test_second_credit_card_not_lost_due_to_address_overlap():
    """Two credit cards separated by addresses should both be detected."""
    text = (
        "Address:\n"
        "17 Lake View Road, Sector 15,\n"
        "Gurugram, Haryana 122001, India\n"
        "Credit Card: 4111 1111 1111 1111\n"
        "Address:\n"
        "28 MG Road, Andheri West,\n"
        "Mumbai, Maharashtra 400058, India\n"
        "Credit Card: 5555 5555 5555 4444"
    )
    detections = detect_pii(text)
    card_dets = [d for d in detections if d['type'] == 'CREDIT_CARD']
    card_texts = [d['text'] for d in card_dets]
    assert any("4111" in t for t in card_texts), f"First card missing: {card_texts}"
    assert any("5555" in t for t in card_texts), f"Second card missing: {card_texts}"


def test_v2_full_synthetic_all_9_categories():
    """Test complete synthetic profile document with all 9 PII categories."""
    text = (
        "Full Name: Aarav Mehta\n"
        "Email: aarav.mehta@example.com\n"
        "Phone: +91 9876543210\n"
        "Date of Birth: 15/08/2001\n"
        "Address:\n"
        "17 Lake View Road, Sector 15,\n"
        "Gurugram, Haryana 122001, India\n"
        "Company:\n"
        "Sharma Technologies Private Limited\n"
        "SSN: 123-45-6789\n"
        "Credit Card: 4111 1111 1111 1111\n"
        "IP Address: 192.168.1.25\n"
        "Name: Neha Kapoor\n"
        "Email: neha.kapoor@example.com\n"
        "Phone: +91 9123456789\n"
        "Address:\n"
        "28 MG Road, Andheri West,\n"
        "Mumbai, Maharashtra 400058, India\n"
        "Credit Card: 5555 5555 5555 4444\n"
    )
    detections = detect_pii(text)
    found_types = {d['type'] for d in detections}
    required = {"PERSON", "EMAIL", "PHONE", "DOB", "ADDRESS", "ORGANIZATION", "SSN", "CREDIT_CARD", "IP_ADDRESS"}
    missing = required - found_types
    assert not missing, f"Missing categories: {missing}"
    person_texts = [d['text'] for d in detections if d['type'] == 'PERSON']
    assert any("Aarav Mehta" in t for t in person_texts), "Aarav Mehta not found"
    assert any("Neha Kapoor" in t for t in person_texts), "Neha Kapoor not found"
    card_texts = [d['text'] for d in detections if d['type'] == 'CREDIT_CARD']
    assert any("4111" in t for t in card_texts), "First card (4111) not detected"
    assert any("5555" in t for t in card_texts), "Second card (5555) not detected"


def test_credit_card_replacement_not_equal_to_original():
    """Ensure fake credit card replacement is different from original and passes Luhn check."""
    redactor = PIIRedactor(seed=42)
    card1 = "4111 1111 1111 1111"
    card2 = "5555 5555 5555 4444"
    rep1 = redactor.get_replacement(card1, "CREDIT_CARD")
    rep2 = redactor.get_replacement(card2, "CREDIT_CARD")

    assert rep1 != card1, f"Replacement was identical to original: {rep1}"
    assert rep2 != card2, f"Replacement was identical to original: {rep2}"
    assert rep1 != rep2, f"Different cards received same replacement: {rep1}"
    assert is_luhn_valid(rep1), f"Fake card 1 is not Luhn valid: {rep1}"
    assert is_luhn_valid(rep2), f"Fake card 2 is not Luhn valid: {rep2}"


def test_repeated_credit_card_and_address_consistency():
    """Check that same card and address get identical replacements when repeated."""
    redactor = PIIRedactor(seed=42)
    card = "4111 1111 1111 1111"
    addr = "17 Lake View Road, Sector 15,\nGurugram, Haryana 122001, India"

    rep_card_1 = redactor.get_replacement(card, "CREDIT_CARD")
    rep_card_2 = redactor.get_replacement(card, "CREDIT_CARD")
    assert rep_card_1 == rep_card_2

    rep_addr_1 = redactor.get_replacement(addr, "ADDRESS")
    rep_addr_2 = redactor.get_replacement(addr, "ADDRESS")
    assert rep_addr_1 == rep_addr_2


def test_credit_card_formats_detection_and_replacement():
    """Test cards with spaces, dashes, or no separators."""
    cards = [
        "4111 1111 1111 1111",
        "4111-1111-1111-1111",
        "4111111111111111",
        "5555 5555 5555 4444",
        "378282246310005"
    ]
    redactor = PIIRedactor(seed=42)
    for c in cards:
        assert is_luhn_valid(c)
        dets = detect_pii(f"My card is {c} here.")
        assert any(d['type'] == 'CREDIT_CARD' for d in dets), f"Card not detected: {c}"
        rep = redactor.get_replacement(c, "CREDIT_CARD")
        assert rep != c, f"Replacement equal to original for {c}"
        assert is_luhn_valid(rep), f"Fake replacement not Luhn valid for {c}: {rep}"


def test_multiline_address_docx_redaction():
    """Test redaction of an address spanning multiple paragraphs in a DOCX."""
    with tempfile.TemporaryDirectory() as tmpdir:
        in_docx = os.path.join(tmpdir, "in.docx")
        out_docx = os.path.join(tmpdir, "out.docx")

        doc = docx.Document()
        doc.add_paragraph("Customer Profile")
        doc.add_paragraph("Address:")
        doc.add_paragraph("17 Lake View Road, Sector 15,")
        doc.add_paragraph("Gurugram, Haryana 122001, India")
        doc.add_paragraph("Company: Acme Private Limited")
        doc.save(in_docx)

        doc = docx.Document(in_docx)
        text_blocks = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n".join(text_blocks)
        detections = detect_pii(full_text)

        redactor = PIIRedactor(seed=42)
        redactor.redact_document(in_docx, out_docx, detections)

        red_doc = docx.Document(out_docx)
        all_paras_text = "\n".join([p.text for p in red_doc.paragraphs])

        assert "17 Lake View Road" not in all_paras_text, "Address line 1 was not redacted"
        assert "Gurugram, Haryana 122001" not in all_paras_text, "Address line 2 was not redacted"
        assert "Maharashtra, India" in all_paras_text or "411 001" in all_paras_text, "Fake address missing"


def test_distinct_fake_generation_for_ssn_dob_ip():
    """Ensure different SSNs, DOBs, and IPs get distinct replacements."""
    redactor = PIIRedactor(seed=42)
    ssn1 = redactor.get_replacement("123-45-6789", "SSN")
    ssn2 = redactor.get_replacement("987-65-4321", "SSN")
    assert ssn1 != ssn2
    assert ssn1 != "123-45-6789"

    dob1 = redactor.get_replacement("15/08/2001", "DOB")
    dob2 = redactor.get_replacement("01/01/1990", "DOB")
    assert dob1 != dob2
    assert dob1 != "15/08/2001"

    ip1 = redactor.get_replacement("192.168.1.25", "IP_ADDRESS")
    ip2 = redactor.get_replacement("10.0.0.50", "IP_ADDRESS")
    assert ip1 != ip2
    assert ip1 != "192.168.1.25"


def test_evaluator_metric_formulas_and_aggregation_consistency():
    """Verify evaluator formulas and metric aggregation across categories."""
    with tempfile.TemporaryDirectory() as tmpdir:
        doc_path = os.path.join(tmpdir, "test_doc.docx")
        gt_path = os.path.join(tmpdir, "gt.json")
        report_path = os.path.join(tmpdir, "report.md")

        doc = docx.Document()
        doc.add_paragraph("The Managing Director is Rashi Patil at Example Technologies Private Limited.")
        doc.add_paragraph("Contact: rashi.patil@example.com, Phone: +91 9876543210.")
        doc.save(doc_path)

        gt_data = {
            "entities": [
                {"type": "PERSON", "text": "Rashi Patil"},
                {"type": "ORGANIZATION", "text": "Example Technologies Private Limited"},
                {"type": "EMAIL", "text": "rashi.patil@example.com"},
                {"type": "PHONE", "text": "+91 9876543210"}
            ]
        }
        with open(gt_path, "w", encoding="utf-8") as f:
            json.dump(gt_data, f)

        res = evaluate_redaction(doc_path, gt_path, report_path)

        # Check that overall sums match category sums
        sum_tp = sum(m["tp"] for m in res["metrics_by_cat"].values())
        sum_fp = sum(m["fp"] for m in res["metrics_by_cat"].values())
        sum_fn = sum(m["fn"] for m in res["metrics_by_cat"].values())

        assert res["tp"] == sum_tp
        assert res["fp"] == sum_fp
        assert res["fn"] == sum_fn

        expected_tn = res["total_tokens"] - res["tp"] - res["fp"] - res["fn"]
        assert res["tn"] == expected_tn

        if res["tp"] + res["fp"] > 0:
            expected_prec = res["tp"] / (res["tp"] + res["fp"])
            assert abs(res["overall_precision"] - expected_prec) < 1e-6
        if res["tp"] + res["fn"] > 0:
            expected_rec = res["tp"] / (res["tp"] + res["fn"])
            assert abs(res["overall_recall"] - expected_rec) < 1e-6
        expected_acc = (res["tp"] + res["tn"]) / res["total_tokens"]
        assert abs(res["overall_accuracy"] - expected_acc) < 1e-6


def test_evaluator_zero_instance_categories_handling():
    """Verify zero-instance categories are marked N/A without causing divide-by-zero errors."""
    with tempfile.TemporaryDirectory() as tmpdir:
        doc_path = os.path.join(tmpdir, "test_doc.docx")
        gt_path = os.path.join(tmpdir, "gt.json")
        report_path = os.path.join(tmpdir, "report.md")

        doc = docx.Document()
        doc.add_paragraph("This text has only one name: Rashi Patil.")
        doc.save(doc_path)

        gt_data = {"entities": [{"type": "PERSON", "text": "Rashi Patil"}]}
        with open(gt_path, "w", encoding="utf-8") as f:
            json.dump(gt_data, f)

        res = evaluate_redaction(doc_path, gt_path, report_path)

        for cat in ["SSN", "CREDIT_CARD", "DOB", "IP_ADDRESS", "ADDRESS"]:
            m = res["metrics_by_cat"][cat]
            assert m["actual_tokens"] == 0
            assert m["precision"] == "N/A"
            assert m["recall"] == "N/A"
            assert m["f1"] == "N/A"
