"""
Handles generating fake replacement data using Faker and replacing text inside DOCX documents.
"""

import os
import re
from faker import Faker
import docx
from typing import Dict, List, Any


class PIIRedactor:
    """Class to manage fake replacements and redact DOCX documents."""

    def __init__(self, seed: int = 42):
        # Use Indian locale and fixed seed so replacements are deterministic
        self.fake = Faker('en_IN')
        Faker.seed(seed)
        self.replacement_map: Dict[str, str] = {}
        self.canonical_person_map: Dict[str, str] = {}
        self.canonical_org_map: Dict[str, str] = {}
        self.canonical_phone_map: Dict[str, str] = {}
        self.canonical_address_map: Dict[str, str] = {}
        self.canonical_card_map: Dict[str, str] = {}
        self.canonical_ssn_map: Dict[str, str] = {}
        self.canonical_dob_map: Dict[str, str] = {}
        self.canonical_ip_map: Dict[str, str] = {}

    def _generate_fake_credit_card(self, original_text: str) -> str:
        """Generates a fake card number with a valid Luhn checksum."""
        clean_digits = re.sub(r'\D', '', original_text)
        num_digits = len(clean_digits) if 13 <= len(clean_digits) <= 19 else 16

        # Pick prefix and generate random digits for the body
        prefix = "5424" if clean_digits.startswith("4") else "4532"
        middle = self.fake.numerify(text="########")
        partial = prefix + middle
        sub = (partial + self.fake.numerify(text="###"))[:num_digits - 1]

        # Calculate Luhn check digit for the last digit
        digits = [int(c) for c in sub]
        checksum = 0
        reverse_digits = digits[::-1]
        for i, digit in enumerate(reverse_digits):
            if i % 2 == 0:
                doubled = digit * 2
                checksum += doubled - 9 if doubled > 9 else doubled
            else:
                checksum += digit
        check_digit = (10 - (checksum % 10)) % 10
        fake_raw = sub + str(check_digit)

        # Match original formatting (spaces or hyphens)
        if " " in original_text:
            chunks = [fake_raw[i:i+4] for i in range(0, len(fake_raw), 4)]
            return " ".join(chunks)
        elif "-" in original_text:
            chunks = [fake_raw[i:i+4] for i in range(0, len(fake_raw), 4)]
            return "-".join(chunks)
        else:
            return fake_raw

    def get_replacement(self, original_text: str, entity_type: str) -> str:
        """Returns a consistent fake replacement for a detected entity."""
        if original_text in self.replacement_map:
            return self.replacement_map[original_text]

        norm_key = original_text.strip().lower()
        replacement = ""

        if entity_type == "PERSON":
            if norm_key not in self.canonical_person_map:
                fake_name = self.fake.name()
                if fake_name == original_text:
                    fake_name = "Aryan Maharaj"
                self.canonical_person_map[norm_key] = fake_name
            fake_base = self.canonical_person_map[norm_key]
            # Match original uppercase styling if applicable
            replacement = fake_base.upper() if original_text.isupper() else fake_base

        elif entity_type == "EMAIL":
            parts = original_text.split('@')
            username = parts[0]
            matched_fake_username = None

            # If email contains a person's first name, try to align the fake email with that fake name
            for person_norm_key, fake_name in self.canonical_person_map.items():
                first_name = person_norm_key.split()[0]
                if first_name and first_name in username.lower():
                    clean_fake = re.sub(r'[^a-zA-Z]', '', fake_name.lower())
                    matched_fake_username = clean_fake
                    break

            if matched_fake_username:
                replacement = f"{matched_fake_username}@example.com"
            else:
                fake_user = re.sub(r'[^a-zA-Z0-9]', '', self.fake.user_name().lower())
                replacement = f"{fake_user}@example.com"

        elif entity_type == "PHONE":
            if norm_key not in self.canonical_phone_map:
                random_digits = self.fake.msisdn()[-8:]
                self.canonical_phone_map[norm_key] = f"+91 98{random_digits}"
            replacement = self.canonical_phone_map[norm_key]

        elif entity_type == "ORGANIZATION":
            if norm_key not in self.canonical_org_map:
                fake_org = f"{self.fake.company()} Private Limited"
                if fake_org == original_text:
                    fake_org = "Sample Enterprises Private Limited"
                self.canonical_org_map[norm_key] = fake_org
            fake_org = self.canonical_org_map[norm_key]
            replacement = fake_org.upper() if original_text.isupper() else fake_org

        elif entity_type == "ADDRESS":
            if norm_key not in self.canonical_address_map:
                street = self.fake.street_name()
                city = self.fake.city()
                bldg = self.fake.building_number()
                fake_addr = f"{bldg}, {street}, {city} – 411 001, Maharashtra, India"
                if fake_addr == original_text:
                    fake_addr = "101, Sample Commercial Complex, MG Road, Pune – 411 001, Maharashtra, India"
                self.canonical_address_map[norm_key] = fake_addr
            replacement = self.canonical_address_map[norm_key]

        elif entity_type == "SSN":
            if norm_key not in self.canonical_ssn_map:
                fake_ssn = self.fake.ssn()
                if fake_ssn == original_text:
                    fake_ssn = "987-65-4321"
                self.canonical_ssn_map[norm_key] = fake_ssn
            replacement = self.canonical_ssn_map[norm_key]

        elif entity_type == "CREDIT_CARD":
            if norm_key not in self.canonical_card_map:
                fake_card = self._generate_fake_credit_card(original_text)
                if fake_card == original_text:
                    fake_card = "5424 1234 5678 9012" if " " in original_text else "5424123456789012"
                self.canonical_card_map[norm_key] = fake_card
            replacement = self.canonical_card_map[norm_key]

        elif entity_type == "IP_ADDRESS":
            if norm_key not in self.canonical_ip_map:
                fake_ip = self.fake.ipv4_private()
                if fake_ip == original_text:
                    fake_ip = "10.0.0.2"
                self.canonical_ip_map[norm_key] = fake_ip
            replacement = self.canonical_ip_map[norm_key]

        elif entity_type == "DOB":
            if norm_key not in self.canonical_dob_map:
                fake_d = self.fake.date_of_birth(minimum_age=22, maximum_age=60)
                if "/" in original_text:
                    fake_dob = fake_d.strftime("%d/%m/%Y")
                elif "-" in original_text:
                    fake_dob = fake_d.strftime("%d-%m-%Y")
                else:
                    fake_dob = fake_d.strftime("%B %d, %Y")
                if fake_dob == original_text:
                    fake_dob = "01/01/1995" if "/" in original_text else "01-01-1995"
                self.canonical_dob_map[norm_key] = fake_dob
            replacement = self.canonical_dob_map[norm_key]

        else:
            replacement = f"[REDACTED_{entity_type}]"

        self.replacement_map[original_text] = replacement
        return replacement

    def build_replacement_map(self, detections: List[Dict[str, Any]]) -> Dict[str, str]:
        """Creates fake mappings for all detected entities, longest strings first."""
        sorted_dets = sorted(detections, key=lambda d: len(d['text']), reverse=True)
        for det in sorted_dets:
            self.get_replacement(det['text'], det['type'])
        return self.replacement_map

    def replace_text_in_paragraph(self, paragraph: Any) -> int:
        """Replaces matched entities in a single paragraph while keeping run formatting."""
        if not paragraph.text or not self.replacement_map:
            return 0

        replacements_made = 0
        full_text = paragraph.text

        # Find items present in this paragraph
        present_items = []
        for orig, fake_val in sorted(self.replacement_map.items(), key=lambda item: len(item[0]), reverse=True):
            if orig in full_text:
                present_items.append((orig, fake_val))

        if not present_items:
            return 0

        # Replace text sequentially
        updated_text = paragraph.text
        for orig, fake_val in present_items:
            if orig in updated_text:
                updated_text = updated_text.replace(orig, fake_val)
                replacements_made += 1

        if updated_text != paragraph.text:
            if paragraph.runs:
                paragraph.runs[0].text = updated_text
                for run in paragraph.runs[1:]:
                    run.text = ""
            else:
                paragraph.text = updated_text

        return replacements_made

    def redact_paragraph_list(self, paragraphs: List[Any]) -> int:
        """Redacts paragraphs, handling single paragraphs and multi-line addresses."""
        if not paragraphs or not self.replacement_map:
            return 0

        replacements_made = 0

        # Check multi-line addresses spanning across paragraphs
        for orig, fake_val in sorted(self.replacement_map.items(), key=lambda item: len(item[0]), reverse=True):
            if "\n" not in orig:
                continue
            orig_lines = [l.strip() for l in orig.split("\n") if l.strip()]
            if len(orig_lines) < 2:
                continue

            k = len(orig_lines)
            i = 0
            while i <= len(paragraphs) - k:
                matched = True
                for j in range(k):
                    if orig_lines[j] not in paragraphs[i + j].text:
                        matched = False
                        break

                if matched:
                    # Replace first paragraph with fake value
                    p_first = paragraphs[i]
                    if p_first.runs:
                        p_first.runs[0].text = p_first.text.replace(orig_lines[0], fake_val)
                        for r in p_first.runs[1:]:
                            r.text = ""
                    else:
                        p_first.text = p_first.text.replace(orig_lines[0], fake_val)

                    # Clear subsequent paragraphs that belonged to the address
                    for j in range(1, k):
                        p_sub = paragraphs[i + j]
                        if p_sub.runs:
                            p_sub.runs[0].text = p_sub.text.replace(orig_lines[j], "").strip()
                            for r in p_sub.runs[1:]:
                                r.text = ""
                        else:
                            p_sub.text = p_sub.text.replace(orig_lines[j], "").strip()

                    replacements_made += 1
                    i += k
                else:
                    i += 1

        # Handle remaining normal single-paragraph replacements
        for p in paragraphs:
            replacements_made += self.replace_text_in_paragraph(p)

        return replacements_made

    def redact_document(self, input_path: str, output_path: str, detections: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Redacts input DOCX file and saves output to output_path."""
        if not os.path.exists(input_path):
            raise FileNotFoundError(f"Input file not found: {input_path}")

        # Build replacement mapping
        self.build_replacement_map(detections)

        doc = docx.Document(input_path)
        total_replacements = 0

        # 1. Main body paragraphs
        total_replacements += self.redact_paragraph_list(doc.paragraphs)

        # 2. Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    total_replacements += self.redact_paragraph_list(cell.paragraphs)

        # 3. Headers and footers
        for section in doc.sections:
            if section.header:
                total_replacements += self.redact_paragraph_list(section.header.paragraphs)
            if section.footer:
                total_replacements += self.redact_paragraph_list(section.footer.paragraphs)

        # Ensure output folder exists and save
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)

        return {
            "output_path": output_path,
            "unique_entities_mapped": len(self.replacement_map),
            "total_replacements_applied": total_replacements
        }
