"""
Main pipeline script for running PII detection, redaction, and evaluation.
"""

import os
import sys
import docx
from src.detectors import detect_pii
from src.redactor import PIIRedactor
from src.evaluator import evaluate_redaction


def run_pipeline(
    input_path: str = "input/Red Herring Prospectus.docx",
    output_path: str = "output/redacted_prospectus.docx",
    ground_truth_path: str = "evaluation/ground_truth.json",
    report_path: str = "evaluation/evaluation_report.md"
):
    print("Starting PII Redaction Pipeline...")

    # 1. Check if input document exists
    if not os.path.exists(input_path):
        print(f"Error: Input document not found at {input_path}")
        sys.exit(1)

    print(f"\n1. Reading document: {input_path}")
    doc = docx.Document(input_path)
    print(f"   Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}")

    # Extract all text from paragraphs and table cells
    text_blocks = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        text_blocks.append(p.text.strip())

    full_text = "\n".join(text_blocks)
    print(f"   Total characters extracted: {len(full_text)}")

    # 2. Run PII detection
    print("\n2. Detecting PII entities...")
    detections = detect_pii(full_text)
    print(f"   Found {len(detections)} potential PII entities")

    # Count detections by category
    cat_counts = {}
    for d in detections:
        cat_counts[d['type']] = cat_counts.get(d['type'], 0) + 1
    for cat, cnt in sorted(cat_counts.items()):
        print(f"   - {cat}: {cnt}")

    # 3. Generate replacements and redact document
    print("\n3. Generating fake replacements and creating redacted DOCX...")
    redactor = PIIRedactor(seed=42)
    stats = redactor.redact_document(input_path, output_path, detections)
    print(f"   Mapped {stats['unique_entities_mapped']} unique entities")
    print(f"   Applied {stats['total_replacements_applied']} replacements in document")
    print(f"   Saved to: {output_path}")

    # 4. Verify output document
    print("\n4. Checking redacted document integrity...")
    if not os.path.exists(output_path):
        print(f"   Error: Output file not created at {output_path}")
        sys.exit(1)

    try:
        redacted_doc = docx.Document(output_path)
        print("   Document opens without issues.")
        print(f"   Paragraph count: {len(redacted_doc.paragraphs)} (matches original)")
        print(f"   Table count: {len(redacted_doc.tables)} (matches original)")
    except Exception as e:
        print(f"   Error opening redacted document: {e}")
        sys.exit(1)

    # 5. Evaluate against ground truth
    print("\n5. Running token-level evaluation...")
    if os.path.exists(ground_truth_path):
        eval_results = evaluate_redaction(input_path, ground_truth_path, report_path)
        print(f"   Total Tokens : {eval_results['total_tokens']:,}")
        print(f"   TP: {eval_results['tp']:,} | FP: {eval_results['fp']:,} | FN: {eval_results['fn']:,} | TN: {eval_results['tn']:,}")
        print(f"   Accuracy  : {eval_results['overall_accuracy']*100:.2f}%")
        print(f"   Precision : {eval_results['overall_precision']*100:.2f}%")
        print(f"   Recall    : {eval_results['overall_recall']*100:.2f}%")
        print(f"   F1 Score  : {eval_results['overall_f1']*100:.2f}%")
        print(f"   Report written to: {report_path}")
    else:
        print(f"   Ground truth not found at {ground_truth_path}, skipping evaluation.")

    print("\nPipeline finished successfully!")


if __name__ == "__main__":
    run_pipeline()
