"""
app.py
Streamlit web application for PII detection and redaction demonstration.
Provides file upload, redaction processing, privacy-safe summary, and redacted DOCX download.
"""

import os
import tempfile
import streamlit as st
import docx

from src.detectors import detect_pii
from src.redactor import PIIRedactor

st.set_page_config(
    page_title="PII Redaction Tool",
    page_icon="🔒",
    layout="centered"
)

st.title("🔒 PII Redaction Tool")
st.markdown("""This application detects and redacts Personally Identifiable Information (PII) from Word documents (`.docx`).
It replaces sensitive data (Names, Emails, Phones, Addresses, Company names, etc.) with realistic fake alternatives.""")

uploaded_file = st.file_uploader("Upload a DOCX Document", type=["docx"])

if uploaded_file is not None:
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_in:
        tmp_in.write(uploaded_file.getvalue())
        input_path = tmp_in.name

    st.info("Document uploaded successfully. Processing PII detection and redaction...")

    # Load and extract text
    doc = docx.Document(input_path)
    text_blocks = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        text_blocks.append(p.text.strip())

    full_text = "\n".join(text_blocks)

    # Detect PII
    detections = detect_pii(full_text)

    # Redact Document
    redactor = PIIRedactor(seed=42)
    output_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
    stats = redactor.redact_document(input_path, output_path, detections)

    # Privacy-safe summary breakdown
    st.subheader("📊 Privacy-Safe Summary")
    st.write(f"**Total Entities Detected**: {len(detections)}")
    st.write(f"**Unique Mappings Generated**: {stats['unique_entities_mapped']}")

    ALL_CATEGORIES = [
        "ADDRESS", "CREDIT_CARD", "DOB", "EMAIL",
        "IP_ADDRESS", "ORGANIZATION", "PERSON", "PHONE", "SSN"
    ]
    cat_counts = {cat: 0 for cat in ALL_CATEGORIES}
    for d in detections:
        if d['type'] in cat_counts:
            cat_counts[d['type']] += 1

    st.markdown("### Category Breakdown")
    for cat in ALL_CATEGORIES:
        st.write(f"- **{cat}**: {cat_counts[cat]} instance(s)")

    # Provide Download
    with open(output_path, "rb") as f:
        bytes_data = f.read()

    st.download_button(
        label="⬇️ Download Redacted DOCX",
        data=bytes_data,
        file_name="redacted_prospectus.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    # Clean up temp files
    os.unlink(input_path)
    os.unlink(output_path)
